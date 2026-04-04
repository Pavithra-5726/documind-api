from fastapi import FastAPI, File, UploadFile, Header, HTTPException, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
import anthropic
import base64
import os
import json
from typing import Optional
import re

app = FastAPI(
    title="AI-Powered Document Analysis & Extraction API",
    description="Extract and analyze content from PDF, DOCX, and image files using Claude AI",
    version="1.0.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

API_KEY = os.environ.get("API_KEY", "hackathon-secret-key-2025")
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")

SUPPORTED_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "image/jpeg": "jpeg",
    "image/jpg": "jpeg",
    "image/png": "png",
    "image/webp": "webp",
    "image/gif": "gif",
}

def verify_api_key(authorization: Optional[str] = Header(None)):
    if not authorization:
        raise HTTPException(status_code=401, detail="Authorization header missing")
    token = authorization.replace("Bearer ", "").replace("ApiKey ", "").strip()
    if token != API_KEY:
        raise HTTPException(status_code=401, detail="Invalid API key")
    return token

@app.get("/")
def root():
    return {
        "service": "AI-Powered Document Analysis & Extraction",
        "status": "running",
        "version": "1.0.0",
        "endpoints": {
            "health": "/health",
            "analyze": "/analyze (POST)",
            "docs": "/docs"
        }
    }

@app.get("/health")
def health():
    return {"status": "healthy", "service": "document-analysis-api"}

@app.post("/analyze")
async def analyze_document(
    file: UploadFile = File(...),
    authorization: Optional[str] = Header(None),
    _: str = Depends(verify_api_key)
):
    content_type = file.content_type or ""

    # Normalize content type
    if file.filename and file.filename.endswith(".pdf"):
        content_type = "application/pdf"
    elif file.filename and (file.filename.endswith(".docx")):
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif file.filename and (file.filename.endswith(".jpg") or file.filename.endswith(".jpeg")):
        content_type = "image/jpeg"
    elif file.filename and file.filename.endswith(".png"):
        content_type = "image/png"

    file_type = SUPPORTED_TYPES.get(content_type)
    if not file_type:
        raise HTTPException(
            status_code=415,
            detail=f"Unsupported file type: {content_type}. Supported: PDF, DOCX, JPEG, PNG, WEBP, GIF"
        )

    file_bytes = await file.read()
    if len(file_bytes) > 20 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="File too large. Maximum size is 20MB")

    b64_data = base64.standard_b64encode(file_bytes).decode("utf-8")

    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

    system_prompt = """You are an expert document analyst and data extraction specialist. 
    When given a document, analyze it thoroughly and extract all meaningful information.
    Always respond with valid JSON only — no markdown, no explanation outside JSON.
    Your JSON must follow this exact structure:
    {
      "document_type": "string (e.g., Invoice, Resume, Contract, Report, Form, Letter, etc.)",
      "summary": "string - 2-4 sentence summary of the document",
      "key_information": {
        "title": "string or null",
        "date": "string or null",
        "author_or_sender": "string or null",
        "recipient": "string or null",
        "subject": "string or null"
      },
      "extracted_data": {
        "entities": ["list of important named entities: people, organizations, locations"],
        "dates": ["list of all dates found"],
        "amounts": ["list of all monetary amounts or numbers"],
        "contact_info": {
          "emails": [],
          "phones": [],
          "addresses": []
        },
        "key_points": ["list of 3-7 key points or highlights from the document"]
      },
      "tables": [
        {
          "title": "table title if any",
          "headers": ["col1", "col2"],
          "rows": [["val1", "val2"]]
        }
      ],
      "sentiment": "positive | negative | neutral | mixed",
      "language": "detected language",
      "confidence_score": 0.95,
      "warnings": ["any issues like low quality, partial content, etc."]
    }
    Extract as much detail as possible. If a field has no data, use null or empty array."""

    try:
        if file_type == "pdf":
            message = client.messages.create(
                model="claude-opus-4-5",
                max_tokens=4096,
                system=system_prompt,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "document",
                                "source": {
                                    "type": "base64",
                                    "media_type": "application/pdf",
                                    "data": b64_data
                                }
                            },
                            {
                                "type": "text",
                                "text": "Analyze this document thoroughly and extract all information. Return only valid JSON."
                            }
                        ]
                    }
                ]
            )
        elif file_type == "docx":
            # For DOCX, extract text first using python-docx
            import io
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            tables_text = ""
            for i, table in enumerate(doc.tables):
                tables_text += f"\n[Table {i+1}]\n"
                for row in table.rows:
                    tables_text += " | ".join([cell.text for cell in row.cells]) + "\n"

            message = client.messages.create(
                model="claude-opus-4-5",
                max_tokens=4096,
                system=system_prompt,
                messages=[
                    {
                        "role": "user",
                        "content": f"Analyze this DOCX document content:\n\n{full_text}\n\n{tables_text}\n\nReturn only valid JSON."
                    }
                ]
            )
        else:
            # Image types
            media_type_map = {
                "jpeg": "image/jpeg",
                "png": "image/png",
                "webp": "image/webp",
                "gif": "image/gif"
            }
            media_type = media_type_map.get(file_type, "image/jpeg")
            message = client.messages.create(
                model="claude-opus-4-5",
                max_tokens=4096,
                system=system_prompt,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image",
                                "source": {
                                    "type": "base64",
                                    "media_type": media_type,
                                    "data": b64_data
                                }
                            },
                            {
                                "type": "text",
                                "text": "Analyze this document image thoroughly and extract all information. Return only valid JSON."
                            }
                        ]
                    }
                ]
            )

        raw_response = message.content[0].text.strip()
        # Clean up any markdown fences if present
        raw_response = re.sub(r"^```json\s*", "", raw_response)
        raw_response = re.sub(r"\s*```$", "", raw_response)

        result = json.loads(raw_response)

        return JSONResponse(content={
            "success": True,
            "filename": file.filename,
            "file_type": file_type.upper(),
            "file_size_bytes": len(file_bytes),
            "analysis": result
        })

    except json.JSONDecodeError:
        return JSONResponse(content={
            "success": True,
            "filename": file.filename,
            "file_type": file_type.upper(),
            "file_size_bytes": len(file_bytes),
            "analysis": {
                "summary": raw_response[:500],
                "raw_text": raw_response,
                "document_type": "Unknown",
                "warnings": ["Could not parse structured JSON from AI response"]
            }
        })
    except anthropic.APIError as e:
        raise HTTPException(status_code=502, detail=f"AI service error: {str(e)}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Processing error: {str(e)}")

@app.post("/analyze/batch")
async def analyze_batch(
    files: list[UploadFile] = File(...),
    _: str = Depends(verify_api_key)
):
    if len(files) > 5:
        raise HTTPException(status_code=400, detail="Maximum 5 files per batch request")

    results = []
    for file in files:
        try:
            result = await analyze_document(file=file, authorization=None, _="skip")
            results.append({"filename": file.filename, "status": "success", "data": result})
        except Exception as e:
            results.append({"filename": file.filename, "status": "error", "error": str(e)})

    return {"success": True, "total": len(files), "results": results}
